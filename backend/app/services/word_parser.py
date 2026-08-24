"""Extreu les dades d'un fitxer .docx de fitxa tècnica.

El document segueix l'estructura estàndard de Farinera Coromina:
- Capçalera: Rev, Data revisió, Data comprovació
- Paràgrafs amb etiqueta + valor
- Taules de paràmetres (fisicoquímiques, reològiques, etc.)
"""
import html
import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table


def _extract_images(doc):
    """Extreu totes les imatges incrustades al document.

    Retorna llista de dicts: {'filename': 'image1.png', 'content_type': 'image/png', 'blob': bytes}
    """
    imatges = []
    seen = set()
    for rel_id, rel in doc.part.rels.items():
        if 'image' not in rel.reltype:
            continue
        try:
            part = rel.target_part
        except Exception:
            continue
        blob = part.blob
        if not blob:
            continue
        # Dedup per contingut (mateixa imatge referenciada més d'un cop)
        sig = (len(blob), blob[:32])
        if sig in seen:
            continue
        seen.add(sig)

        partname = getattr(part, 'partname', '') or ''
        original = os.path.basename(str(partname)) or f'image_{len(imatges) + 1}'
        content_type = getattr(part, 'content_type', '') or 'image/png'
        ext = os.path.splitext(original)[1]
        if not ext:
            mime_to_ext = {
                'image/png': '.png',
                'image/jpeg': '.jpg',
                'image/gif': '.gif',
                'image/webp': '.webp',
                'image/bmp': '.bmp',
                'image/svg+xml': '.svg',
            }
            ext = mime_to_ext.get(content_type, '.png')
            original = f'{original}{ext}'

        imatges.append({
            'filename': original,
            'content_type': content_type,
            'blob': blob,
        })
    return imatges


def _clean(text):
    """Neteja text: treure espais duplicats i salts innecessaris."""
    if not text:
        return ''
    return ' '.join(text.strip().split())


def _extract_bilingual(text):
    """Retorna el text complet (bilingüe). Per identificar etiquetes,
    retorna la part castellana per fer match."""
    return text.strip()


def _extract_bilingual_for_match(text):
    """Extreu la part castellana per fer match amb etiquetes."""
    if ' / ' in text:
        return text.split(' / ')[0].strip()
    return text.strip()


# Data en formats dd/mm/aaaa, d-m-aa... (evita agafar trossos d'hores o codis)
_DATE_RE = re.compile(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})')


def _parse_header(doc):
    """Extreu rev, data_revisio, data_comprovacio i el títol de la capçalera.
    Busca tant al header de Word com a les taules del body.

    Algunes fitxes no etiqueten la fila de comprovació (posen "Fecha/Data:" i
    "Fecha/Data rev.:" en comptes de "Fecha/Data Rev:" i "Fecha/Data Comprov.:").
    Per a aquests casos hi ha un fallback posicional: si no s'ha trobat cap
    etiqueta "Comprov.", l'última fila amb data de la capçalera és la data de
    comprovació i la primera la de revisió.
    """
    info = {'rev': '', 'data_revisio': '', 'data_comprovacio': '', 'titol': ''}
    dates_per_fila = []   # una data com a molt per fila, en ordre de document

    def _check_cell(text):
        # Rev.: número — excloent "Fecha/Data Rev: dd/mm/aaaa" que també conté "Rev:"
        if ('Rev.:' in text or 'Rev:' in text) and 'Fecha' not in text and 'Data' not in text:
            m = re.search(r'Rev\.?:\s*(\d+)', text)
            if m:
                info['rev'] = m.group(1)
        low = text.lower()
        if ('fecha' in low or 'data' in low) and 'rev' in low and 'comprov' not in low:
            m = _DATE_RE.search(text)
            if m:
                info['data_revisio'] = m.group(1)
        if 'comprov' in low:
            m = _DATE_RE.search(text)
            if m:
                info['data_comprovacio'] = m.group(1)
        # Títol de la fitxa: serveix per detectar si el document és bilingüe
        if not info['titol'] and ('ficha t' in low or 'fitxa t' in low):
            info['titol'] = text

    def _scan_table(table):
        for row in table.rows:
            data_fila = ''
            for cell in row.cells:
                text = cell.text.strip()
                if not text:
                    continue
                _check_cell(text)
                if not data_fila:
                    m = _DATE_RE.search(text)
                    if m:
                        data_fila = m.group(1)
            if data_fila:
                dates_per_fila.append(data_fila)

    # Buscar al header de Word
    for section in doc.sections:
        header = section.header
        if not header or not header.tables:
            continue
        for table in header.tables:
            _scan_table(table)
        break

    # Si no s'ha trobat, buscar a les primeres taules del body
    if not info['rev']:
        for table in doc.tables[:3]:
            _scan_table(table)
            if info['rev']:
                break

    # Fallback posicional quan la capçalera no etiqueta la data de comprovació
    if not info['data_comprovacio'] and len(dates_per_fila) >= 2:
        info['data_comprovacio'] = dates_per_fila[-1]
        if not info['data_revisio']:
            info['data_revisio'] = dates_per_fila[0]

    return info


def _cell_parametre_html(cell):
    """Extreu el text d'una cel·la de paràmetre preservant els paràgrafs i
    marcant els que estan en cursiva-gris al Word (peus regulatoris com
    'Según RD 677/2016') amb el span de text secundari."""
    parts = []
    for p in cell.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        parts.append(_wrap_secondary(text, _is_secondary_paragraph(p)))
    return '<br>'.join(parts)


def _parse_param_table(table):
    """Extreu files parametre/valor d'una taula estàndard.
    Manté el text complet bilingüe dels paràmetres i notes."""
    rows = []
    for row in table.rows:
        cells_param = [_cell_parametre_html(cell) for cell in row.cells]
        cells_plain = [cell.text.strip() for cell in row.cells]
        # Dedup merged cells (comparant text pla)
        deduped_param = []
        deduped_valor = []
        prev = None
        for cp, ct in zip(cells_param, cells_plain):
            if ct != prev:
                deduped_param.append(cp)
                deduped_valor.append(ct)
            prev = ct

        if len(deduped_param) < 2:
            continue

        param = deduped_param[0]
        valor = deduped_valor[1]

        # Saltar capçaleres de taula (comparem sobre text pla, sense HTML)
        param_plain = deduped_valor[0].lower().strip()
        skip = ['parámetro', 'parámetro / paràmetre', 'valor', 'valor límite']
        if param_plain in skip:
            continue
        if not param or deduped_valor[0] == valor:
            continue

        rows.append({'parametre': param, 'valor': valor})

    return rows


# Mapeig d'etiquetes de paràgrafs a camps del contingut
FIELD_MAP = {
    'código de referencia': 'codi_referencia',
    'codi de referència': 'codi_referencia',
    'certificación': 'certificacio',
    'certificació': 'certificacio',
    'denominación comercial del producto': 'denominacio_comercial',
    'denominació comercial del producte': 'denominacio_comercial',
    'denominación jurídica del producto': 'denominacio_juridica',
    'denominació jurídica del producte': 'denominacio_juridica',
    'código ean': 'codi_ean',
    'codi ean': 'codi_ean',
    'descripción del producto': 'descripcio',
    'descripció del producte': 'descripcio',
    'origen del producto y procedencia del cereal': 'origen',
    'origen del producte i procedència del cereal': 'origen',
    'ingredientes': 'ingredients',
    'ingredients': 'ingredients',
    'alérgenos': 'alergens',
    'al·lèrgens': 'alergens',
    'ogm': 'ogm',
    'irradiación – ionización': 'irradiacio',
    'irradiación': 'irradiacio',
    'irradiació': 'irradiacio',
    'características organolépticas': 'caract_organoleptiques',
    'característiques organolèptiques': 'caract_organoleptiques',
    'presentación – envase': 'presentacio_envase',
    'presentació': 'presentacio_envase',
    'uso previsto': 'us_previst',
    'ús previst': 'us_previst',
    'condiciones de almacenaje': 'condicions_emmagatzematge',
    "condicions d'emmagatzematge": 'condicions_emmagatzematge',
    'condiciones de transporte': 'condicions_transport',
    'condicions de transport': 'condicions_transport',
    'vida útil del producto': 'vida_util',
    'vida útil del producte': 'vida_util',
    'otra legislación aplicable': 'legislacio_aplicable',
    'altra legislació aplicable': 'legislacio_aplicable',
    'producto fabricado para (razón social)': 'fabricat_per',
    'producto fabricado para': 'fabricat_per',
    'producte fabricat per a': 'fabricat_per',
    'vigencia del documento': 'vigencia_document',
    'vigència del document': 'vigencia_document',
    'pesticidas': 'pesticidas',
    'pesticides': 'pesticidas',
}

# Títols de seccions a ignorar (no són camps de text)
SECTION_TITLES = {
    'características físico – químicas',
    'características físico-químicas',
    'características reológicas',
    'características microbiológicas',
    'características higiénico-sanitarias',
    'características higiénico – sanitarias',
    'características higiénico sanitarias',
    'parámetros de contaminantes',
    'valores nutricionales',
    'pesticidas',
}

# Prefixos de títol de secció. Serveixen de reserva quan el Word usa una
# variant no catalogada a SECTION_TITLES (p.ex. "Características Higiénico-
# Sanitarias"): sense això el títol s'acabaria guardant com a peu de la taula
# anterior. Els camps reals (Características organolépticas) no hi arriben mai
# perquè _match_field_label s'avalua abans.
SECTION_TITLE_PREFIXES = (
    'características', 'característiques',
    'parámetros de', 'paràmetres de',
    'valores nutricionales', 'valors nutricionals',
)

# Títols de taules i el camp JSON corresponent
TABLE_MAP = {
    'humedad': 'fisicoquimiques',
    'proteína': 'fisicoquimiques',
    'w': 'reologiques',
    'p/l': 'reologiques',
    'aerobios': 'microbiologiques',
    'parámetros microbiológicos': 'microbiologiques',
    'micotoxinas': 'micotoxines',
    'aflatoxina': 'micotoxines',
    'alcaloides': 'alcaloides',
    'metales pesados': 'metalls_pesants',
    'cadmio': 'metalls_pesants',
    'pesticidas': 'pesticidas_taula',
    'valores nutricionales': 'valors_nutricionals',
    'valor energético': 'valors_nutricionals',
}


def _identify_table(table):
    """Identifica a quina secció pertany una taula pel seu contingut."""
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip().lower()
            text_clean = _extract_bilingual_for_match(text).lower()
            for key, field in TABLE_MAP.items():
                if key in text_clean or key in text:
                    return field
    return None


# Estil per text secundari (peu regulatori). Ha de coincidir amb SECONDARY_STYLE
# de frontend/src/components/FitxaForm.jsx per garantir la paritat visual.
SECONDARY_STYLE = 'font-size: 0.75em; color: #595959; font-style: italic;'

# Tokens que identifiquen un character style equivalent a "Subtle Emphasis"
# (cursiva + gris). Cobreix ÉnfasisSutil (ES), SubtleEmphasis (EN), i variants
# retornades per python-docx (ex: 'nfasissutil' amb la lletra inicial no ASCII).
_SUBTLE_EMPHASIS_TOKENS = ('emphasis', 'subtle', 'sutil', 'nfasissutil')


def _run_is_secondary(rPr):
    """True si el run té formatació de text secundari (cursiva + gris)."""
    if rPr is None:
        return False
    rStyle = rPr.find(qn('w:rStyle'))
    if rStyle is not None:
        style_id = (rStyle.get(qn('w:val')) or '').lower()
        if any(tok in style_id for tok in _SUBTLE_EMPHASIS_TOKENS):
            return True
    if rPr.find(qn('w:i')) is not None:
        return True
    color = rPr.find(qn('w:color'))
    if color is not None:
        val = (color.get(qn('w:val')) or '').lower().lstrip('#')
        if len(val) == 6:
            try:
                r, g, b = int(val[0:2], 16), int(val[2:4], 16), int(val[4:6], 16)
            except ValueError:
                return False
            if abs(r - g) <= 8 and abs(g - b) <= 8 and 0x30 <= r <= 0xA0:
                return True
    return False


def _is_secondary_paragraph(p):
    """True si tots els runs amb text no buit d'un paràgraf estan en cursiva-gris."""
    p_elem = p._element if hasattr(p, '_element') else p
    total = 0
    secondary = 0
    for run in p_elem.findall(qn('w:r')):
        text_parts = run.findall(qn('w:t'))
        run_text = ''.join(t.text or '' for t in text_parts)
        if not run_text.strip():
            continue
        total += 1
        if _run_is_secondary(run.find(qn('w:rPr'))):
            secondary += 1
    return total > 0 and secondary == total


def _wrap_secondary(text, is_secondary):
    """Embolcalla el text amb el span de text secundari si escau, escapant HTML."""
    if not is_secondary:
        return text
    return f'<span style="{SECONDARY_STYLE}">{html.escape(text, quote=False)}</span>'


def _iter_body_items(doc):
    """Itera els elements del body en ordre de document: (kind, obj)
    on kind és 'p' (paràgraf) o 't' (taula)."""
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn('w:p'):
            yield ('p', Paragraph(child, doc))
        elif tag == qn('w:tbl'):
            yield ('t', Table(child, doc))


def _match_field_label(text):
    """Retorna el nom del camp si el text coincideix amb una etiqueta, o None."""
    text_for_match = _extract_bilingual_for_match(text).lower().rstrip('.')
    for label, field_name in FIELD_MAP.items():
        if text_for_match.startswith(label) or label.startswith(text_for_match):
            return field_name
    return None


def _is_section_title(text):
    text_for_match = _extract_bilingual_for_match(text).lower().rstrip('.')
    if text_for_match in SECTION_TITLES:
        return True
    # Reserva acotada: títol curt, sense xifres, amb prefix conegut
    if len(text_for_match) < 60 and not any(ch.isdigit() for ch in text_for_match):
        return text_for_match.startswith(SECTION_TITLE_PREFIXES)
    return False


def _detect_idioma(titol, etiquetes_total, etiquetes_bilingues):
    """Retorna 'es' (només castellà) o 'bilingue' (castellà / català).

    Es decideix per la proporció d'etiquetes escrites en els dos idiomes. Si el
    document no en té prou de reconegudes, es mira el títol de la capçalera
    ("FICHA TÉCNICA" sol vs "FICHA TÉCNICA / FITXA TÈCNICA").
    """
    if etiquetes_total >= 4:
        return 'es' if (etiquetes_bilingues / etiquetes_total) < 0.5 else 'bilingue'
    if titol and ' / ' not in titol:
        return 'es'
    return 'bilingue'


def parse_docx(file_path):
    """Parseja un fitxer .docx de fitxa tècnica i retorna un dict amb les dades.

    Returns:
        dict amb:
            - 'contingut': dict amb tots els camps
            - 'rev': número de revisió
            - 'data_revisio': data de revisió
            - 'data_comprovacio': data de comprovació
            - 'art_codi': codi de referència (per identificar la fitxa)
            - 'imatges': llista de dicts {'filename', 'content_type', 'blob'} extretes del docx
    """
    doc = Document(file_path)

    # 1. Capçalera
    header_info = _parse_header(doc)

    # 2. Iterar en ordre de document: paràgrafs + taules barrejats
    # Inicialitzem tots els _note a buit per garantir que en re-importar sobre
    # una fitxa existent els peus vells no es queden a la BD si el docx nou no
    # els inclou.
    contingut = {
        'fisicoquimiques_note': '',
        'reologiques_note': '',
        'microbiologiques_note': '',
        'micotoxines_note': '',
        'alcaloides_note': '',
        'metalls_pesants_note': '',
        'valors_nutricionals_note': '',
    }
    tables_data = {
        'fisicoquimiques': [],
        'reologiques': [],
        'microbiologiques': [],
        'micotoxines': [],
        'alcaloides': [],
        'metalls_pesants': [],
        'valors_nutricionals': [],
    }
    current_field = None       # camp de text actiu (per paràgrafs següents)
    last_table_key = None      # última taula processada (per capturar el peu)
    etiquetes_total = 0        # etiquetes reconegudes (per detectar l'idioma)
    etiquetes_bilingues = 0    # ...de les quals escrites "castellà / català"

    for kind, item in _iter_body_items(doc):
        if kind == 'p':
            text = item.text.strip()
            if not text:
                continue  # línia en blanc: NO resetejar current_field

            # Nova etiqueta → canvi de camp
            field = _match_field_label(text)
            if field is not None:
                etiquetes_total += 1
                if ' / ' in text:
                    etiquetes_bilingues += 1
                current_field = field
                last_table_key = None
                continue

            # Títol de secció → reset (però mantenir last_table_key per notes)
            if _is_section_title(text):
                current_field = None
                continue

            # Text lliure: prioritzar peu de taula si acabem de processar una taula.
            # Els peus s'emmagatzemen com a text pla (el frontend/PDF els estilitza
            # ja com a secundari per convenció).
            if last_table_key and current_field is None:
                note_key = f'{last_table_key}_note'
                existing = contingut.get(note_key, '')
                contingut[note_key] = (existing + '\n' + text) if existing else text
                continue

            # Camps de text lliure: preservar format secundari del Word amb span
            if current_field:
                wrapped = _wrap_secondary(text, _is_secondary_paragraph(item))
                existing = contingut.get(current_field, '')
                contingut[current_field] = (existing + '\n' + wrapped) if existing else wrapped

        elif kind == 't':
            table_type = _identify_table(item)
            current_field = None  # una taula tanca el camp de text actiu

            if table_type == 'pesticidas_taula':
                # Taula amb estructura títol + text (fila 0 = "Pesticidas / Pesticides",
                # files següents = contingut bilingüe). No és paràmetre/valor.
                # Preserva format secundari (EnfasisSutil al Word) de cada paràgraf.
                text_parts = []
                seen = set()
                for row in item.rows[1:]:
                    for cell in row.cells:
                        cell_html = _cell_parametre_html(cell)
                        plain = cell.text.strip()
                        if cell_html and plain not in seen:
                            text_parts.append(cell_html)
                            seen.add(plain)
                if text_parts and not contingut.get('pesticidas'):
                    contingut['pesticidas'] = '<br>'.join(text_parts)
                last_table_key = None
                continue

            if table_type and table_type in tables_data:
                rows = _parse_param_table(item)
                tables_data[table_type].extend(rows)
                last_table_key = table_type
            else:
                last_table_key = None

    contingut.update(tables_data)

    # Idioma del document: les fitxes de client (PBUK...) són només en castellà
    # i no han de portar els títols duplicats castellà/català.
    contingut['_idioma'] = _detect_idioma(
        header_info.get('titol', ''), etiquetes_total, etiquetes_bilingues)

    # Extreure art_codi
    art_codi = contingut.get('codi_referencia', '').strip()

    # Extreure imatges incrustades
    imatges = _extract_images(doc)

    return {
        'contingut': contingut,
        'rev': header_info['rev'],
        'data_revisio': header_info['data_revisio'],
        'data_comprovacio': header_info['data_comprovacio'],
        'art_codi': art_codi,
        'imatges': imatges,
    }
